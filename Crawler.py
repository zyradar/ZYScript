import re
import requests
from bs4 import BeautifulSoup
from docx import Document
import pandas as pd
import numpy as np
import cv2
import tkinter as tk
from tkinter import messagebox
from GetParameters import get_windowstate


class CRawler:
    def __init__(self):
        # 指定要爬取的网站URL
        # self.url = "https://baike.baidu.com/item/%E5%BE%B7%E7%8E%9B%E8%A5%BF%E4%BA%9A%E4%B9%8B%E5%8A%9B/7087378"
        self.url = 'https://pic.quanjing.com/28/h3/QJ5100545083.jpg@%21350h'
        # self.url = 'https://zh.wikipedia.org/wiki/%E7%9B%96%E4%BC%A6'
        # self.url = "https://www.bing.com/ck/a?!&&p=634fd1d6c790ba04c39705640bffd3abae8091eca1a3c2e16baff36587c7e44cJmltdHM9MTc0MDg3MzYwMA&ptn=3&ver=2&hsh=4&fclid=1971f260-452d-6218-1dca-e7c444b363f8&psq=%e7%9b%96%e4%bc%a6&u=a1aHR0cHM6Ly96aC53aWtpcGVkaWEub3JnL3dpa2kvJUU3JTlCJTk2JUU0JUJDJUE2&ntb=1"
        self.save_path = 'C:/Users/HZY/Pictures/work/txt/'
        self.save_name = 'test'
        self.save_format = 'jpg'
        self.format = {'txt': '.txt', 'word': '.docx', 'excel': '.xlsx'}
        self.text_search = None
        self.search_number = 5
        self.search_Engines = {'百度': 'https://www.baidu.com/s?wd=',
                               '搜狗': 'https://www.sogou.com/web?query=',
                               '必应': 'https://www.bing.com/search?q=',
                               '360': 'https://www.so.com/s?q=',
                               '神马': 'https://m.sm.cn/s?q=',
                               '谷歌': 'https://www.google.com/search?q='}
        self.engine_name = '必应'
        self.top_results = []
        # "User-Agent"	用户代理（User-Agent）用于标识请求是从什么设备、浏览器和操作系统发出的。
        # "Mozilla/5.0"	表示这是 Mozilla 浏览器的一个版本。
        # "Windows NT 10.0; Win64; x64"	表示使用 Windows 10 64 位系统。
        # "AppleWebKit/537.36"	表示使用 WebKit 537.36 作为浏览器内核（常见于 Chrome 和 Safari）。
        # "like Gecko"	兼容 Gecko 内核（Mozilla 的浏览器引擎）。
        # "Chrome/110.0.0.0"	表示使用的是 Chrome 110.0.0.0 版本。
        # "Safari/537.36"	表示兼容 Safari 浏览器的 WebKit 537.36 版本。
        # 这个 User-Agent 表示 请求来自 Chrome 110.0 浏览器，运行在 Windows 10 64 位操作系统上。
        self.browser = 'edge'
        self.proxy = '127.0.0.1:7890'
        self.proxies = {                                # netsh winhttp show proxy  查看代理服务器
            'http': 'http://' + self.proxy,            # HTTP 代理
            'https': 'http://' + self.proxy            # HTTPS 代理
        }
        self.user_agents = {
            'edge': "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko)"
                    " Chrome/120.0.0.0 Safari/537.36 Edg/120.0.0.0",
            'lenovo': "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko)"
                      " Chrome/110.0.0.0 LenovoBrowser/23.0.0.0 Safari/537.36",
            'chrome': "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko)"
                      " Chrome/110.0.0.0 Safari/537.36"
        }
        self.headers = {
            'User-Agent': self.user_agents.get(self.browser, self.user_agents[self.browser])
            # "User-Agent": random.choice(list(user_agents.values()))
        }
        self.open_window = False
        self.show_time = 1

    def grawler_text(self, flag):
        self.top_results.clear()
        self.set_save_file()
        if flag != '0':
            self.url = self.search_Engines[self.engine_name] + flag
            # self.url = f"https://www.bing.com/search?q={flag}"
        self.headers = {
            'User-Agent': self.user_agents.get(self.browser, self.user_agents[self.browser])
            # "User-Agent": random.choice(list(user_agents.values()))
        }
        if self.proxy:
            self.proxies = {                                # netsh winhttp show proxy  查看代理服务器
                'http': 'http://' + self.proxy,             # HTTP 代理
                'https': 'http://' + self.proxy             # HTTPS 代理
            }
            # 发送 GET 请求到 self.url，并附带 self.headers 模拟浏览器访问。
            response = requests.get(self.url, headers=self.headers, proxies=self.proxies)
        else:
            response = requests.get(self.url, headers=self.headers)
        # 检查请求是否成功，200 表示成功获取网页内容。404 表示网页不存在。403 表示被拒绝访问（可能是爬虫被拦截）。
        if response.status_code == 200:
            soup = BeautifulSoup(response.text, "html.parser")                              # 解析HTML
            if flag != '0':
                # results = soup.select("li.b_algo h2 a")         # 提取网站链接
                results = soup.select("li.b_algo")              # 获取每个搜索结果的完整信息
                for result in results:
                    title_tag = result.select_one("h2 a")  # 标题和链接
                    snippet = result.select_one(".b_caption p")  # 摘要信息
                    if title_tag and snippet:
                        title = title_tag.text
                        link = title_tag["href"]
                        description = snippet.text.strip()
                        self.top_results.append((title, link, description))
                # **这里可以加一个“筛选访问量最多”步骤，比如按可信网站（如 Wikipedia, 网易）排序**
                self.top_results = sorted(self.top_results, key=lambda x: len(x[2]), reverse=True)[:self.search_number]  # 按摘要长度筛选
                # for title, link, description in self.top_results:    # 标题（title），链接（link），摘要（description）
                #     print(title, link, description)
                if self.top_results:
                    self.choice_result()
                else:
                    messagebox.showwarning('警告', '未能正常爬取信息，开发者正在寻找原因，尽量尽快测试并修复此报错！！')
                self.url = self.text_search
                # if self.proxy:
                #     response = requests.get(self.url, headers=self.headers, proxies=self.proxies, allow_redirects=True)
                # else:
                #     response = requests.get(self.url, headers=self.headers, allow_redirects=True)
                # soup = BeautifulSoup(response.text, 'html.parser')          # 假设您已经爬取了包含重定向URL的页面
                # redirect_match = re.search(r'window\.location\.replace\("(https://.*)"\);', soup.prettify())    # 使用正则表达式查找跳转到实际链接的URL
                # print('redirect_match', redirect_match)
                # if redirect_match:
                #     final_url = redirect_match.group(1)
                #     print("最终链接：", final_url)
                #     response = requests.get(final_url)              # 获取最终页面的内容
                #     content = response.text
                #
                # print("response.url", response.url)
                # self.url = response.url        # 获取最终的URL
                # print("url", self.url)
                # if response.status_code == 200:
                #     soup = BeautifulSoup(response.text, "html.parser")  # 解析HTML
                # else:
                #     messagebox.showwarning('警告', f"{response.status_code}")
            text = soup.get_text()                                                          # 获取所有文本内容
            # text = str(self.top_results)
            if self.save_format == 'word':
                doc = Document()                                                            # 创建一个新的 Word 文档
                doc.add_paragraph(text)                                                     # 添加文本内容
                doc.save(self.save_path + self.save_name + self.format[self.save_format])   # 保存为指定路径的 .docx 文件
                if flag != '0':
                    linkdoc = Document()
                    for title, link, description in self.top_results:
                        linkdoc.add_paragraph('标题：' + title + '\n' + '链接：' + link +
                                              '\n' + '摘要：' + description + '\n\n\n')
                    linkdoc.save(self.save_path + self.save_name + '相关链接' + self.format[self.save_format])
            if self.save_format == 'txt':
                with open(self.save_path + self.save_name + self.format[self.save_format], 'a', encoding="utf-8") as f:
                    f.writelines(text)
                if flag != '0':
                    with open(self.save_path + self.save_name + '相关链接' + self.format[self.save_format], 'a',
                              encoding="utf-8") as f:
                        for title, link, description in self.top_results:
                            # f.writelines(str(self.top_results))
                            f.write(title + '---' + link + '---' + description + '\n')
            if self.save_format == 'excel':
                df = pd.DataFrame([{"URL": self.url, "Content": text}])                     # 转换为 DataFrame 封装为字典
                # df = pd.DataFrame({'Content': text.split("\n")})                          # 拆分为多行存储
                df.to_excel(self.save_path + self.save_name + self.format[self.save_format], index=False, engine='openpyxl')
        else:
            messagebox.showwarning('警告', f"{response.status_code}")
        messagebox.showinfo('温馨提示', "grawler_text任务已完成")

    def grawler_image(self, flag):
        self.open_window, self.show_time = get_windowstate()
        self.set_save_file()
        self.headers = {
            'User-Agent': self.user_agents.get(self.browser, self.user_agents[self.browser])
            # "User-Agent": random.choice(list(user_agents.values()))
        }
        if self.proxy:
            self.proxies = {                                # netsh winhttp show proxy  查看代理服务器
                'http': 'http://' + self.proxy,             # HTTP 代理
                'https': 'http://' + self.proxy             # HTTPS 代理
            }
            image_bytes = np.asarray(
                bytearray(requests.get(self.url, headers=self.headers, proxies=self.proxies).content),
                dtype=np.uint8)
        else:
            image_bytes = np.asarray(
                bytearray(requests.get(self.url, headers=self.headers).content),
                dtype=np.uint8)
        img = cv2.imdecode(image_bytes, cv2.IMREAD_COLOR)
        img = cv2.resize(img, (1920, 1200))
        cv2.imwrite(self.save_path + self.save_name + '.' + self.save_format, img)
        if self.open_window:
            self.take_gui(img)
        messagebox.showinfo('温馨提示', "grawler_image任务已完成")

    def set_save_file(self):
        set_window = tk.Tk()
        set_window.geometry("750x600")
        set_window.title("爬取信息保存设置")
        tk.Label(set_window, text="因涉及本地路径，该窗口配置不可跳过！", font=("Arial", 18)).grid(row=0, column=0, pady=10)
        tk.Label(set_window, text="爬取的网站地址(若使用文本搜索引擎，可不填):", font=("Arial", 16)).grid(row=1, column=0, pady=10)
        url_entry = tk.Entry(set_window, width=30)
        url_entry.grid(row=1, column=1, pady=5)
        tk.Label(set_window, text="请选择需要伪装的浏览器，如'chrome':", font=("Arial", 16)).grid(row=2, column=0, pady=5)
        browser_entry = tk.Entry(set_window, width=30)
        browser_entry.grid(row=2, column=1, pady=5)
        tk.Label(set_window, text="设置爬取的热门信息待选数(默认为5,文本搜索专用):", font=("Arial", 16)).grid(row=3, column=0, pady=5)
        search_number_entry = tk.Entry(set_window, width=30)
        search_number_entry.grid(row=3, column=1, pady=5)
        tk.Label(set_window, text="本地代理IP(若本地未启用代理，可不填):", font=("Arial", 16)).grid(row=4, column=0, pady=5)
        proxy_entry = tk.Entry(set_window, width=30)
        proxy_entry.grid(row=4, column=1, pady=5)
        tk.Label(set_window, text="设置文件保存的文件夹路径:", font=("Arial", 16)).grid(row=5, column=0, pady=5)
        save_path_entry = tk.Entry(set_window, width=30)
        save_path_entry.grid(row=5, column=1, pady=5)
        tk.Label(set_window, text="设置文件保存的名字:", font=("Arial", 16)).grid(row=6, column=0, pady=5)
        save_name_entry = tk.Entry(set_window, width=30)
        save_name_entry.grid(row=6, column=1, pady=5)
        tk.Label(set_window, text="设置保存的文件类型，如'word':", font=("Arial", 16)).grid(row=7, column=0, pady=5)
        save_format_entry = tk.Entry(set_window, width=30)
        save_format_entry.grid(row=7, column=1, pady=5)

        def on_submit():                                    # 创建提交按钮
            self.url = url_entry.get() if url_entry.get() else self.url
            self.browser = browser_entry.get() if browser_entry.get() else self.browser
            self.search_number = search_number_entry.get() if search_number_entry.get() else self.search_number
            self.proxy = proxy_entry.get() if proxy_entry.get() else None
            self.save_path = save_path_entry.get() if save_path_entry.get() else self.save_path
            self.save_name = save_name_entry.get() if save_name_entry.get() else self.save_name
            self.save_format = save_format_entry.get() if save_format_entry.get() else self.save_format
            set_window.destroy()                          # 关闭输入窗口
        tk.Button(set_window, text="提交", command=on_submit, width=20, bg="lightblue", font=("Arial", 16)).grid(row=8, column=0, columnspan=2, pady=10)
        tk.Label(set_window, text="当前系统内部配置伪装浏览器仅有:\nchrome，lenovo，edge\n仅支持爬取信息保存文件类型为:\n"
                                  "word,txt,excel以及一切图片格式（jpg,png等）", font=("Arial", 16)).grid(row=9, column=0, pady=10)
        set_window.wait_window()                          # 等待直到输入窗口关闭

    def choice_result(self):
        set_window = tk.Tk()
        set_window.geometry("700x500")
        set_window.title("请选择您要爬取的信息")
        i = 0

        def on_submit(t, l):
            self.text_search = l
            messagebox.showwarning('警告！！！！', '想什么呢，出bug了，developer睡大觉去了，没人修bug。'
                                   '关键字爬取的内容已经保存到你设置的路径去了，至于你刚刚选择并点击的要爬取内容，developer不帮你爬，'
                                   'developer已经把链接帮你保存下来了，你要是想爬那些选定的内容，自己去复制链接，'
                                   '去浏览器打开链接，在浏览器的标签页上再复制一遍浏览器上的链接，在此工具下手动输入爬取网址进行内容爬取，'
                                   '只有文本搜索模块出现bug了，手动输入网址爬取模块仍能正常运行')
            set_window.destroy()                          # 关闭输入窗口
        for title, link, description in self.top_results:    # 标题（title），链接（link），摘要（description）
            tk.Label(set_window, text=title, font=("Arial", 12)).grid(row=i, column=0, pady=10)
            tk.Button(set_window, text="爬取", command=lambda t=title, l=link: on_submit(t, l), width=20, bg="lightblue", font=("Arial", 12)).grid(row=i, column=1, columnspan=1, pady=10)
            i += 1
        tk.Label(set_window, text="已为你爬取到访问量最高,最热点的内容\n请选取其中任一内容进行爬取", font=("Arial", 16)).grid(row=i, column=0, pady=40)
        set_window.wait_window()                          # 等待直到输入窗口关闭

    def take_gui(self, img):
        cv2.namedWindow("gui", cv2.WINDOW_NORMAL)
        cv2.imshow("gui", img)
        if cv2.waitKey(self.show_time) == 27:
            self.open_window = False
            cv2.destroyAllWindows()


# CRawler().grawler_text('螨虫')
# CRawler().set_save_file()
# 127.0.0.1:7890
